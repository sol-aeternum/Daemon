from __future__ import annotations

import os
import tempfile
from pathlib import Path
from unittest.mock import patch

import pytest

from orchestrator.subagents.base import SubagentType
from orchestrator.subagents.document import DocumentSubagent
from orchestrator.tools.spawn import _persist_file_result


class TestDocumentIntegration:
    @pytest.mark.asyncio
    async def test_spawn_document_subagent_wiring(self):
        config = {
            "openrouter_api_key": "test-key",
            "openrouter_base_url": "https://openrouter.ai/v1",
        }

        agent = DocumentSubagent(config)

        assert agent.agent_type == SubagentType.DOCUMENT
        assert agent.model == "anthropic/claude-sonnet-4.5"

    def test_file_persistence_flow(self):
        result_dict = {
            "agent_type": "document",
            "data": {
                "format": "docx",
                "filename": "example.docx",
                "file_url": "/generated-files/example.docx",
                "generation_code": "print('docx')",
            },
            "metadata": {},
        }

        processed = _persist_file_result(result_dict)

        assert "generation_code" in processed["metadata"]
        assert processed["metadata"]["generation_code"] == "print('docx')"

    def test_csv_file_persistence_flow(self):
        result_dict = {
            "agent_type": "document",
            "data": {
                "format": "csv",
                "filename": "example.csv",
                "file_url": "/generated-files/example.csv",
                "generation_code": "import csv",
            },
            "metadata": {},
        }

        processed = _persist_file_result(result_dict)

        assert "generation_code" in processed["metadata"]
        assert processed["metadata"]["generation_code"] == "import csv"

    @pytest.mark.asyncio
    async def test_document_execution_with_generation_code(self):
        config = {
            "openrouter_api_key": "test-key",
            "openrouter_base_url": "https://openrouter.ai/v1",
        }
        agent = DocumentSubagent(config)

        mock_code = (
            "from docx import Document\n"
            "doc = Document()\n"
            "doc.add_paragraph('Hello World')\n"
            "doc.save('output.docx')\n"
        )

        with patch(
            "orchestrator.subagents.document.load_document_skill",
            return_value="Skill content",
        ):
            with patch.object(
                agent,
                "_generate_code",
                return_value={"success": True, "code": mock_code},
            ):
                with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
                    _ = f.write(b"PK\x03\x04test")
                    temp_path = f.name

                try:
                    with patch.object(
                        agent,
                        "_execute_sandbox",
                        return_value={"success": True, "file_path": temp_path},
                    ):
                        with patch.object(agent, "_persist_file", return_value=Path(temp_path)):
                            result = await agent.execute(
                                "Create a simple document", {"format": "docx"}
                            )

                            assert result.success is True
                            assert "generation_code" in result.data
                            assert result.data["generation_code"] == mock_code
                finally:
                    if os.path.exists(temp_path):
                        os.unlink(temp_path)

    @pytest.mark.asyncio
    async def test_full_generation_flow_with_revision(self):
        config = {
            "openrouter_api_key": "test-key",
            "openrouter_base_url": "https://openrouter.ai/v1",
        }
        agent = DocumentSubagent(config)

        initial_code = (
            "from docx import Document\n"
            "doc = Document()\n"
            "doc.add_paragraph('Version 1')\n"
            "doc.save('output.docx')\n"
        )
        revision_code = (
            "from docx import Document\n"
            "doc = Document()\n"
            "doc.add_paragraph('Version 2 - Revised')\n"
            "doc.save('output.docx')\n"
        )

        with patch(
            "orchestrator.subagents.document.load_document_skill",
            return_value="Skill content",
        ):
            with patch.object(
                agent,
                "_generate_code",
                side_effect=[
                    {"success": True, "code": initial_code},
                    {"success": True, "code": revision_code},
                ],
            ):
                with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
                    _ = f.write(b"PK\x03\x04test")
                    temp_path = f.name

                try:
                    with patch.object(
                        agent,
                        "_execute_sandbox",
                        return_value={"success": True, "file_path": temp_path},
                    ):
                        with patch.object(agent, "_persist_file", return_value=Path(temp_path)):
                            result1 = await agent.execute("Create a document", {"format": "docx"})
                            assert result1.success is True

                            result2 = await agent.execute(
                                "Make it say Version 2",
                                {
                                    "format": "docx",
                                    "generation_code": result1.data["generation_code"],
                                },
                            )
                            assert result2.success is True
                            assert result2.data["generation_code"] == revision_code
                finally:
                    if os.path.exists(temp_path):
                        os.unlink(temp_path)


class TestDocumentEndpointIntegration:
    def test_generated_files_path_semantics(self):
        generated_files_dir = Path("data/generated_files")
        assert generated_files_dir.name == "generated_files"
        assert generated_files_dir.parent.name == "data"


class TestDocumentSkillLoading:
    def test_load_document_skill_raises_for_invalid_format(self):
        from orchestrator.skills_loader import load_document_skill

        with pytest.raises(ValueError, match="Document skill for format"):
            _ = load_document_skill("pptx")

    def test_load_document_skill_supports_docx(self):
        from orchestrator.skills_loader import load_document_skill

        with patch.object(Path, "exists", return_value=True):
            with patch.object(
                Path,
                "read_text",
                return_value="---\nname: docx\n---\nUse python-docx to build documents.",
            ):
                skill = load_document_skill("docx")
                assert skill == "Use python-docx to build documents."

    def test_load_document_skill_supports_csv(self):
        from orchestrator.skills_loader import load_document_skill

        with patch.object(Path, "exists", return_value=True):
            with patch.object(
                Path,
                "read_text",
                return_value="---\nname: csv\n---\nUse Python csv module to write rows.",
            ):
                skill = load_document_skill("csv")
                assert skill == "Use Python csv module to write rows."

    def test_load_document_skill_allows_13kb_without_truncation(self):
        from orchestrator.skills_loader import load_document_skill

        body = "a" * 13000
        content = f"---\nname: docx\n---\n{body}"

        with patch.object(Path, "exists", return_value=True):
            with patch.object(Path, "read_text", return_value=content):
                skill = load_document_skill("docx")

        assert len(skill) == 13000
        assert skill == body

    def test_load_document_skill_truncates_17kb_with_warning(
        self, caplog: pytest.LogCaptureFixture
    ):
        from orchestrator.skills_loader import (
            SUBAGENT_SKILL_CHAR_LIMIT,
            load_document_skill,
        )

        body = "b" * 17000
        content = f"---\nname: docx\n---\n{body}"

        with patch.object(Path, "exists", return_value=True):
            with patch.object(Path, "read_text", return_value=content):
                with caplog.at_level("WARNING"):
                    skill = load_document_skill("docx")

        assert len(skill) == SUBAGENT_SKILL_CHAR_LIMIT
        assert "exceeded" in caplog.text
        assert "truncating" in caplog.text

    def test_orchestrator_skill_injection_limit_remains_2000(self):
        from orchestrator.skills_store import build_enabled_skills_block

        long_content = "x" * 2500
        skill_summary = {
            "id": "demo-skill",
            "name": "demo-skill",
            "description": "demo",
            "enabled": True,
            "updated_at": "2026-03-09T00:00:00Z",
        }
        skill_detail = {
            "id": "demo-skill",
            "name": "demo-skill",
            "description": "demo",
            "enabled": True,
            "updated_at": "2026-03-09T00:00:00Z",
            "content": long_content,
        }

        with patch("orchestrator.skills_store.list_skills", return_value=[skill_summary]):
            with patch("orchestrator.skills_store.get_skill", return_value=skill_detail):
                block = build_enabled_skills_block()

        assert ("x" * 2000) in block
        assert ("x" * 2001) not in block
