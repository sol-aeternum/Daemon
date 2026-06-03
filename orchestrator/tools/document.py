"""Deterministic document generation tool - CSV and DOCX."""

from __future__ import annotations

import csv
import json
import logging
import uuid
from pathlib import Path
from typing import Any

from orchestrator.tools.registry import Tool
from orchestrator.utils import slugify_filename

logger = logging.getLogger(__name__)

GENERATED_FILES_DIR = Path(__file__).resolve().parent.parent.parent / "data" / "generated_files"


class GenerateDocumentTool(Tool):
    name = "generate_document"
    description = "Generate a document file (CSV or DOCX) from structured content"
    parameters = {
        "type": "object",
        "properties": {
            "format": {
                "type": "string",
                "description": "Document format: 'csv' for spreadsheet or 'docx' for Word document",
                "enum": ["csv", "docx"],
            },
            "content": {
                "type": "string",
                "description": "Text content for the document body, or CSV rows as JSON string (list of lists with header row). For CSV, can be a JSON array of arrays or a header+rows object.",
            },
            "title": {
                "type": "string",
                "description": "Optional document title",
            },
            "filename": {
                "type": "string",
                "description": "Optional short kebab-case filename without extension (e.g. 'quarterly-report-2026'). If not provided, a random name is generated.",
            },
            "sections": {
                "type": "array",
                "description": "Optional list of section titles and body text, e.g. [{'title': 'Introduction', 'body': '...'}, ...]",
                "items": {
                    "type": "object",
                    "properties": {
                        "title": {"type": "string"},
                        "body": {"type": "string"},
                    },
                },
            },
            "table": {
                "type": "object",
                "description": "Optional table data with 'headers' (list of strings) and 'rows' (list of lists). For CSV, this is an alternative to passing rows via content.",
                "properties": {
                    "headers": {"type": "array", "items": {"type": "string"}},
                    "rows": {"type": "array", "items": {"type": "array"}},
                },
            },
        },
        "required": ["format", "content"],
    }

    async def execute(
        self,
        format: str,
        content: str,
        title: str = "",
        filename: str = "",
        sections: list[dict[str, str]] | None = None,
        table: dict[str, Any] | None = None,
        **kwargs: Any,
    ) -> str:
        doc_format = format.lower()
        if doc_format not in ("csv", "docx"):
            return json.dumps({
                "success": False,
                "error": f"Unsupported format: {format}. Supported: csv, docx.",
            })

        try:
            if doc_format == "csv":
                result = await self._generate_csv(content, table, title, filename)
            else:
                result = await self._generate_docx(content, sections, table, title, filename)
            return json.dumps(result)
        except Exception as e:
            logger.exception("Document generation failed")
            return json.dumps({"success": False, "error": str(e)})

    async def _generate_csv(
        self,
        content: str,
        table: dict[str, Any] | None,
        title: str,
        filename: str,
    ) -> dict[str, Any]:
        GENERATED_FILES_DIR.mkdir(parents=True, exist_ok=True)

        rows: list[list[str]] = []
        headers: list[str] = []

        # Try to parse content as JSON rows
        try:
            parsed = json.loads(content)
            if isinstance(parsed, list) and all(isinstance(r, list) for r in parsed):
                if parsed:
                    headers = [str(h) for h in (parsed[0] if isinstance(parsed[0], list) else [])]
                    rows = [[str(c) for c in r] for r in parsed]
            elif isinstance(parsed, dict) and "headers" in parsed and "rows" in parsed:
                headers = [str(h) for h in parsed["headers"]]
                rows = [[str(c) for c in r] for r in parsed["rows"]]
        except (json.JSONDecodeError, TypeError):
            pass

        # Fall back to text content as single column
        if not rows:
            headers = ["Content"]
            for line in content.strip().splitlines():
                if line.strip():
                    rows.append([line.strip()])

        safe_ext = "csv"
        slug = slugify_filename(filename) if filename else ""

        if slug:
            out_name = f"{slug}.{safe_ext}"
            out_path = GENERATED_FILES_DIR / out_name
            if out_path.exists():
                slug = f"{slug}-{uuid.uuid4().hex[:8]}"
                out_name = f"{slug}.{safe_ext}"
                out_path = GENERATED_FILES_DIR / out_name
        else:
            while True:
                out_name = f"{uuid.uuid4().hex[:12]}.{safe_ext}"
                out_path = GENERATED_FILES_DIR / out_name
                if not out_path.exists():
                    break

        with out_path.open("w", newline="", encoding="utf-8") as fh:
            writer = csv.writer(fh)
            if headers:
                writer.writerow(headers)
            writer.writerows(rows)

        file_size = out_path.stat().st_size
        return {
            "success": True,
            "data": {
                "file_url": f"/generated-files/{out_name}",
                "filename": out_name,
                "format": "csv",
                "file_size": file_size,
                "mime_type": "text/csv",
            },
        }

    async def _generate_docx(
        self,
        content: str,
        sections: list[dict[str, str]] | None,
        table: dict[str, Any] | None,
        title: str,
        filename: str,
    ) -> dict[str, Any]:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        GENERATED_FILES_DIR.mkdir(parents=True, exist_ok=True)

        doc = Document()

        if title:
            heading = doc.add_heading(title, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if content.strip():
            doc.add_paragraph(content.strip())

        if sections:
            for section in sections:
                sec_title = section.get("title", "")
                sec_body = section.get("body", "")
                if sec_title:
                    doc.add_heading(sec_title, level=2)
                if sec_body:
                    doc.add_paragraph(sec_body.strip())

        if table:
            headers = table.get("headers", [])
            rows_data = table.get("rows", [])
            if headers and rows_data:
                tbl = doc.add_table(rows=1, cols=len(headers))
                tbl.style = "Table Grid"
                hdr_cells = tbl.rows[0].cells
                for i, h in enumerate(headers):
                    hdr_cells[i].text = str(h)
                    for para in hdr_cells[i].paragraphs:
                        for run in para.runs:
                            run.bold = True
                for row_data in rows_data:
                    row_cells = tbl.add_row().cells
                    for i, cell_val in enumerate(row_data):
                        if i < len(row_cells):
                            row_cells[i].text = str(cell_val)

        safe_ext = "docx"
        slug = slugify_filename(filename) if filename else ""

        if slug:
            out_name = f"{slug}.{safe_ext}"
            out_path = GENERATED_FILES_DIR / out_name
            if out_path.exists():
                slug = f"{slug}-{uuid.uuid4().hex[:8]}"
                out_name = f"{slug}.{safe_ext}"
                out_path = GENERATED_FILES_DIR / out_name
        else:
            while True:
                out_name = f"{uuid.uuid4().hex[:12]}.{safe_ext}"
                out_path = GENERATED_FILES_DIR / out_name
                if not out_path.exists():
                    break

        doc.save(str(out_path))

        file_size = out_path.stat().st_size
        return {
            "success": True,
            "data": {
                "file_url": f"/generated-files/{out_name}",
                "filename": out_name,
                "format": "docx",
                "file_size": file_size,
                "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            },
        }
